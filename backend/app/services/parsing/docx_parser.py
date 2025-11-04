"""
DOCX Parser - DOCX/DOC 파일에서 텍스트 추출 (개선된 라이브러리 사용)
"""
import logging
from typing import Dict, Any
import re

logger = logging.getLogger(__name__)

class DOCXParser:
    """DOCX/DOC 파일 파서 - 개선된 라이브러리 사용"""
    
    def extract_text(self, file_path: str) -> str:
        """
        DOCX 파일에서 텍스트 추출 (PDF 변환 후 파싱)
        
        Args:
            file_path: DOCX 파일 경로
            
        Returns:
            추출된 텍스트
        """
        logger = logging.getLogger(__name__)
        
        try:
            # 방법 1: DOCX를 PDF로 변환 후 파싱
            try:
                import tempfile
                import os
                from app.services.parsing.pdf_parser import PDFParser
                
                # 임시 PDF 파일 생성
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                    temp_pdf_path = temp_pdf.name
                
                # DOCX를 PDF로 변환
                self._convert_docx_to_pdf(file_path, temp_pdf_path)
                
                # PDF 파싱
                pdf_parser = PDFParser()
                text = pdf_parser.extract_text(temp_pdf_path)
                
                # 임시 파일 삭제
                os.unlink(temp_pdf_path)
                
                logger.info(f"DOCX → PDF 변환 후 파싱 완료: {len(text)} 문자")
                return text
                
            except Exception as e:
                logger.warning(f"DOCX → PDF 변환 실패: {e}. 직접 파싱으로 재시도")
            
            # 방법 2: 직접 DOCX 파싱 (기존 방법)
            from docx import Document
            doc = Document(file_path)
            text = ""
            
            logger.info("DOCX 직접 파싱 시작")
            logger.info(f"총 단락 수: {len(doc.paragraphs)}")
            logger.info(f"총 표 수: {len(doc.tables)}")
            
            # 모든 요소를 순서대로 파싱 (단락과 표를 번갈아가며)
            all_elements = []
            
            # 단락 추가
            for i, paragraph in enumerate(doc.paragraphs):
                all_elements.append(('paragraph', i, paragraph))
            
            # 표 추가
            for i, table in enumerate(doc.tables):
                all_elements.append(('table', i, table))
            
            # 요소들을 문서 순서대로 정렬 (실제로는 단락과 표가 섞여있을 수 있음)
            # 단락을 먼저 처리하고 표를 나중에 처리
            for element_type, idx, element in all_elements:
                if element_type == 'paragraph':
                    paragraph_text = element.text.strip()
                    if paragraph_text:
                        text += paragraph_text + "\n"
                        logger.info(f"단락 {idx+1}: '{paragraph_text[:50]}...' (길이: {len(paragraph_text)})")
                elif element_type == 'table':
                    logger.info(f"표 {idx+1} 파싱 시작: {len(element.rows)}행, {len(element.columns)}열")
                    
                    table_text = ""
                    for row_idx, row in enumerate(element.rows):
                        row_text = ""
                        for cell_idx, cell in enumerate(row.cells):
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text += cell_text + " | "
                                logger.info(f"  셀 [{row_idx+1},{cell_idx+1}]: {cell_text[:30]}...")
                        if row_text:
                            table_text += row_text + "\n"
                    
                    if table_text.strip():
                        text += f"\n[표 {idx+1}]\n" + table_text + "\n"
                        logger.info(f"표 {idx+1} 파싱 완료: {len(table_text)} 문자")
            
            # 텍스트 정리
            text = self._clean_text(text)
            logger.info(f"DOCX 직접 파싱 완료: {len(text)} 문자")
            return text
            
        except Exception as e:
            logger.error(f"DOCX 텍스트 추출 실패: {e}")
            raise Exception(f"DOCX 텍스트 추출 실패: {e}")
    
    def _convert_docx_to_pdf(self, docx_path: str, pdf_path: str) -> None:
        """
        DOCX 파일을 PDF로 변환
        
        Args:
            docx_path: DOCX 파일 경로
            pdf_path: 출력 PDF 파일 경로
        """
        try:
            # LibreOffice 사용 (시스템에 설치되어 있는 경우)
            import subprocess
            import os
            
            # LibreOffice 명령어로 변환
            output_dir = os.path.dirname(pdf_path)
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                docx_path
            ]
            
            logger.info(f"LibreOffice 변환 명령어: {' '.join(cmd)}")
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            
            logger.info(f"LibreOffice 변환 결과: returncode={result.returncode}")
            logger.info(f"LibreOffice stdout: {result.stdout}")
            logger.info(f"LibreOffice stderr: {result.stderr}")
            
            if result.returncode == 0:
                # 변환된 파일명 찾기
                base_name = os.path.splitext(os.path.basename(docx_path))[0]
                expected_pdf = os.path.join(output_dir, f"{base_name}.pdf")
                
                logger.info(f"예상 PDF 파일 경로: {expected_pdf}")
                logger.info(f"출력 디렉토리 내용: {os.listdir(output_dir)}")
                
                if os.path.exists(expected_pdf):
                    os.rename(expected_pdf, pdf_path)
                    logger.info(f"LibreOffice로 DOCX → PDF 변환 완료: {docx_path} → {pdf_path}")
                    return
                else:
                    raise Exception(f"변환된 PDF 파일을 찾을 수 없음: {expected_pdf}")
            else:
                raise Exception(f"LibreOffice 변환 실패: {result.stderr}")
                
        except Exception as e:
            logger.error(f"DOCX → PDF 변환 실패: {e}")
            raise Exception(f"DOCX → PDF 변환 실패: {e}")
    
    def _clean_text(self, text: str) -> str:
        """텍스트 정리"""
        # 불필요한 공백 제거
        text = text.strip()
        
        # 연속된 공백을 하나로
        import re
        text = re.sub(r' +', ' ', text)
        
        # 연속된 줄바꿈을 두 개로
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        return text
    
    def parse_resume(self, text: str) -> Dict[str, Any]:
        """
        이력서 텍스트에서 구조화된 정보 추출
        (PDF Parser와 동일한 로직 사용)
        """
        from app.services.parsing.pdf_parser import PDFParser
        pdf_parser = PDFParser()
        return pdf_parser.parse_resume(text)


